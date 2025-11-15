# import os
# from PyPDF2 import PdfReader
# try:
#     import docx
# except Exception:
#     docx = None

# def extract_text(file_path):
#     text = ''
#     lower = file_path.lower()
#     if lower.endswith('.pdf'):
#         try:
#             reader = PdfReader(file_path)
#             for page in reader.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text + "\n"
#         except Exception as e:
#             text = ''
#     elif lower.endswith('.docx') and docx is not None:
#         try:
#             d = docx.Document(file_path)
#             text = "\n".join([p.text for p in d.paragraphs])
#         except Exception:
#             text = ''
#     else:
#         # fallback: try reading as text
#         try:
#             with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
#                 text = f.read()
#         except Exception:
#             text = ''
#     return text


import os
from PyPDF2 import PdfReader

try:
    import docx
except ImportError:
    docx = None


def extract_text(file_path):
    text = ""
    lower = file_path.lower()

    # ------------------------------
    # PDF FILES
    # ------------------------------
    if lower.endswith(".pdf"):
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
        except Exception as e:
            print("PDF extraction failed:", e)
            return ""

    # ------------------------------
    # DOCX FILES
    # ------------------------------
    elif lower.endswith(".docx") and docx is not None:
        try:
            document = docx.Document(file_path)
            paragraphs = [p.text for p in document.paragraphs if p.text]
            text = "\n".join(paragraphs)
        except Exception as e:
            print("DOCX extraction failed:", e)
            return ""

    # ------------------------------
    # Fallback for TXT, unknown formats
    # ------------------------------
    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            print("Fallback extraction failed:", e)
            return ""

    return text
